import os.path, re, docx, pandas as pd
from os import listdir
from os.path import isfile, join
from tqdm.auto import tqdm
from time import gmtime, strftime

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = './Planilha/'
final = './Gerados/'

file_success = []
file_with_errors = []
list_errors = []

list_doc = []

proccess = False
commands = {
	1: 'Presidente',
	2: 'Conscrito',
	3: 'Medico',
	4: 'Odonto'
}

print('1 - Gerar Declaração')
print('2 - Gerar Entrevista')
print('3 - Gerar Inspeção Médica')
print('4 - Gerar Inspeção Odontológica')

def set_font_default(document):
	font = document.styles['Normal'].font
	font.name = 'Times New Roman'

def get_header(document):
	p = document.add_paragraph()
	p.add_run('\n\n')
	logo = document.add_picture('./__assets/images/logo-top.jpg', width=Inches(1))
	logo_paragraph = document.paragraphs[-1] 
	logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	title = document.add_paragraph().add_run('MINISTÉRIO DA DEFESA\nEXÉRCITO BRASILEIRO\n\nBASE DE ADMINISTRAÇÃO E APOIO DO IBIRAPUERA\n2º REGIÃO MILITAR DO SULDESTE')
	title.bold = True
	title_paragraph = document.paragraphs[-1] 
	title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def get_subtitle(document, subtitle):
	subtitle = document.add_paragraph().add_run(subtitle)
	subtitle.bold = True
	subtitle.underline = True
	subtitle_paragraph = document.paragraphs[-1] 
	subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def model_1(name, ra, date, hour):
	try:

		dt = str(date).split('-')
		hour = str(hour).split(':')
		now = strftime("%d/%m/%Y", gmtime())

		strdate = '/'.join([dt[2].split(' ')[0], dt[1], dt[0]])
		strhour = 'h'.join([hour[0], hour[1]])
		strhourleft = 'h'.join([str(int(hour[0])+4), hour[1]])

		document = Document()

		set_font_default(document)
		get_header(document)
		get_subtitle(document, 'D E C L A R A Ç Ã O')

		paragraph = document.add_paragraph().add_run(f'\nDeclaro, para fins de justificativa de faltas, de acordo com o art. 195 do Regulamento da Lei do Serviço Militar, que o Conscrito {name}, de Certificado de Alistamento Militar nº {str(ra)}, compareceu à Seleção Complementar no dia {strdate}, das {strhour} até {strhourleft}, no Base de Administração e Apoio do Ibirapuera, situado a Rua Manoel da Nóbrega, nº 1015, Paraíso, São Paulo – SP, com a finalidade de cumprir suas obrigações legais atinentes ao Serviço Militar.')
		paragraph_paragraph = document.paragraphs[-1] 
		paragraph_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

		locale = document.add_paragraph().add_run(f'\nSão Paulo-SP, {str(now)}.\n\n')
		locale_paragraph = document.paragraphs[-1] 
		locale_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

		assign = document.add_paragraph().add_run('WERNER ALVES SILVEIRA - Major\nPresidente da Seleção Complementar')
		assign.alignment = WD_ALIGN_PARAGRAPH.CENTER

		assign_paragraph = document.paragraphs[-1] 
		assign_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

		namefile = re.sub(u'[^a-zA-Z0-9áéíóúÁÉÍÓÚâêîôÂÊÎÔãõÃÕçÇ: ]', '', name)
		namefile = re.sub(r'\s', '-', namefile)

		document.save(f'./{final}/Declaração/{namefile.upper()}.docx')
		file_success.append(namefile)

	except Exception as e:
		print(f'\n{e}')

def model_2(item):
	try:

		now = strftime("%d/%m/%Y", gmtime())

		document = Document()

		set_font_default(document)
		get_header(document)
		get_subtitle(document, 'FICHA DE ENTREVISTA DO CONSCRITO')

		interviewer = document.add_paragraph().add_run(f'ENTREVISTADOR: {item[97]}')
		interviewer.bold = True

		p = document.add_paragraph().add_run('1. IDENTIFICAÇÃO')
		p.bold = True

		document.add_paragraph().add_run('a. DADOS GERAIS')
		paragraph = document.add_paragraph(f'1) NOME COMPLETO: {item[1]}\n2) DATA DE NASCIMENTO: -\n3) LOCAL DE NASCIMENTO: -\n4) IDENTIDADE:  -\n5) CPF: -\n6) NOME DO PAI: -\n7) NOME DA MÃE: -\n8) TIPO SANGUÍNEO/FATOR Rh: -\n9) TÍTULO ELEITORAL:       Zona:       Sessão: \n10) CNH/CAT: {item[17]}\n11) ESCOLARIDADE: {item[18]}')
		paragraph.paragraph_format.left_indent = Inches(0.25)

		document.add_paragraph().add_run('b. ENDEREÇO')
		paragraph = document.add_paragraph(f'1) LOGRADOURO: {item[2]}\n2) NÚMERO: {item[3]}\n3) COMPLEMENTO: {item[4]} \n4) BAIRRO: {item[5]}\n\n5) PONTO DE REFERÊNCIA: {item[6]}\n6) ZONA: {item[7]}\n7) CIDADE: -\n8) CEP: {item[8]}\n9) TELEFONE PESSOAL: -\n10) TELEFONE PARA RECADOS: {item[9]}\n11) E-MAIL: \n12) FACEBOOK: {item[10]}\n13) INSTAGRAM: {item[11]}\n14) TWITTER: {item[12]}')
		paragraph.paragraph_format.left_indent = Inches(0.25)

		p = document.add_paragraph().add_run('\n2. ATIVIDADES')
		p.bold = True

		paragraph = document.add_paragraph(f'a. FIRMA QUE TRABALHA: {item[20]}\nb. FUNÇÃO: {item[21]}\nc. ESTABELECIMENTO DE ENSINO QUE ESTUDA: {item[22]}\nd. ESPORTES QUE PRATICA: {item[23]}\ne. CLUBES QUE FREQUENTA: {item[24]}\nf. INSTRUMENTOS MUSICAIS QUE TOCA: {item[25]}\ng. HABILIDADES QUE POSSUI: -')
		paragraph.paragraph_format.left_indent = Inches(0.25)

		p = document.add_paragraph().add_run('\n3. PSICOSSOCIAL')
		p.bold = True

		paragraph = document.add_paragraph(f'a. JÁ EXPERIMENTOU ALGUM TIPO DE DROGA: \nb. CONSOME ALGUM TIPO DE DROGA: \nc. PARTICIPA DE JOGOS DE AZAR: \nd. MOVIMENTOS SOCIAIS: \ne. MOVIMENTOS POLÍTICOS: \nf. MOVIMENTOS RELIGIOSOS: \ng. JÁ SE ENVOLVEU COM ALGUM ÓRGÃO DE SEGURANÇA PÚBLICA, MESMO NA CONDIÇÃO DE TESTEMUNHA? . \nh. ALGUÉM DA FAMÍLIA JÁ SE ENVOLVEU COM ALGUM ÓRGÃO DE SEGURANÇA PÚBLICA, MESMO NA CONDIÇÃO DE TESTEMUNHA? . \ni. NAS PROXIMIDADES DO LOCAL ONDE RESIDE EXISTEM PROBLEMAS COM RELAÇÃO A TRÁFICO OU USO DE DROGAS? . \nj. POSSUI AÇÕES NA JUSTIÇA CONTRA O ESTADO OU A NÍVEL FEDERAL? . ')
		paragraph.paragraph_format.left_indent = Inches(0.25)

		namefile = re.sub(u'[^a-zA-Z0-9áéíóúÁÉÍÓÚâêîôÂÊÎÔãõÃÕçÇ: ]', '', item[1])
		namefile = re.sub(r'\s', '-', namefile)
		document.save(f'./{final}/Entrevistas/{namefile.upper()}.docx')
		file_success.append(namefile)

	except Exception as e:
		print(f'\n{e}')

def load_module(url, sheet, cols):
	try:
		d = pd.read_excel(url, sheet_name=sheet, usecols=cols)
		df = pd.DataFrame(data=d)

		for item in df.values:

			if sheet == 'Presidente' and str(item[4].lower()) == 'sim': 
				model_1(item[0], item[1], item[2], item[3])

			if sheet == 'Conscrito':
				model_2(item)
				break

	except Exception as e:
		list_errors.append({'error': e, 'file': url})

files = [f for f in listdir(path) if isfile(join(path, f))]

while not proccess:

	try:

		load = int(input('\n>> Escolha dentre as opções acima: '))
		proccess = commands[load] if commands[load] != None else commands[1]

		if (len(files) > 0):
			
			for f in files:

				if load == 1:
					print('\nPor favor aguarde, gerando declarações...')
					load_module(f'{path}{f}', 'Presidente', range(0, 30))

				if load == 2:
					print('\nPor favor aguarde, gerando declarações...')
					load_module(f'{path}{f}', 'Conscrito', range(0, 99))

			if len(list_errors) == 0: 
				print(f'Foram gerados ({len(file_success)}) declarações com sucesso.\n')
				proccess = False

		else:
			print('\nNão existem arquivos no diretório.')
			break
	except:
		print('\nComando inválido. Tente novamente.')
		proccess = False

if (len(list_errors) > 0):
	print('\nForam encontrados estes erros:')
	for i in list_errors: 
		msg = i['error']
		filename = i['file']
		print(f'>> {msg}, "{filename}"')

input('\nPressione qualquer tecla para continuar...')